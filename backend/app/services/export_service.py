from html import escape
from io import BytesIO

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.schemas.report import ReportContent, ReportDetail, ReportEditorConfig, ReportSection


class ExportService:
    """报告派生产物导出服务。

    content_json 是权威数据；HTML 只是派生产物，供预览/导出使用。
    """

    def render_report_html(self, report: ReportDetail) -> str:
        content = report.content_json or report.draft_json
        if not content:
            return self._html_document(report.title, "<p class=\"empty\">暂无报告内容。</p>", report.editor_config)
        body = self._render_content(content)
        return self._html_document(content.title or report.title, body, report.editor_config)

    def render_report_docx(self, report: ReportDetail) -> bytes:
        content = report.content_json or report.draft_json
        document = Document()
        self._configure_docx_page(document, report.editor_config)
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)

        if report.html_snapshot:
            self._append_html_to_docx(document, report.html_snapshot)
        elif content:
            for section in content.sections:
                if section.type == "html":
                    self._append_html_to_docx(document, section.content or "")
                else:
                    document.add_heading(section.title, level=2)
                    if section.content:
                        document.add_paragraph(section.content)
        else:
            document.add_heading(report.title, level=1)
            document.add_paragraph("暂无报告内容。")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _configure_docx_page(self, document: Document, editor_config: ReportEditorConfig) -> None:
        page = editor_config.page
        section = document.sections[0]
        if page.size and page.size.width and page.size.height:
            width = Cm(page.size.width)
            height = Cm(page.size.height)
            if page.orientation == "landscape":
                width, height = max(width, height), min(width, height)
            else:
                width, height = min(width, height), max(width, height)
            section.page_width = width
            section.page_height = height
        elif page.orientation == "landscape":
            section.page_width, section.page_height = section.page_height, section.page_width
        if page.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
        section.top_margin = Cm(page.margin.top)
        section.bottom_margin = Cm(page.margin.bottom)
        section.left_margin = Cm(page.margin.left)
        section.right_margin = Cm(page.margin.right)

    def _append_html_to_docx(self, document: Document, html: str) -> None:
        """Convert the editor HTML while retaining its visible formatting.

        UMO writes font, size, colour, alignment, spacing and similar formatting
        into inline HTML styles. The old exporter used get_text(), which discarded
        all of those styles. Keep the HTML snapshot as the source for Word runs.
        """
        soup = BeautifulSoup(html, "html.parser")
        root = soup.body or soup
        for node in root.children:
            if isinstance(node, NavigableString):
                if str(node).strip():
                    document.add_paragraph(str(node))
            elif isinstance(node, Tag):
                self._append_html_node(document, node)

    def _append_html_node(self, document: Document, node: Tag) -> None:
        name = node.name.lower()
        if name in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "blockquote", "pre"}:
            level = min(int(name[1]), 6) if name.startswith("h") else None
            paragraph = document.add_heading("", level=level) if level else document.add_paragraph()
            if name == "blockquote":
                paragraph.paragraph_format.left_indent = Cm(0.74)
            self._apply_paragraph_style(paragraph, node)
            self._append_inline_content(paragraph, node, heading_level=level)
            return
        if name in {"ul", "ol"}:
            list_style = "List Bullet" if name == "ul" else "List Number"
            for item in node.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style=list_style)
                self._apply_paragraph_style(paragraph, item)
                self._append_inline_content(paragraph, item)
            return
        if name == "table":
            rows = node.find_all("tr")
            column_count = max((len(row.find_all(["th", "td"], recursive=False)) for row in rows), default=0)
            if not rows or not column_count:
                return
            table = document.add_table(rows=len(rows), cols=column_count)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                cells = row.find_all(["th", "td"], recursive=False)
                for column_index, html_cell in enumerate(cells):
                    cell = table.cell(row_index, column_index)
                    paragraph = cell.paragraphs[0]
                    self._apply_paragraph_style(paragraph, html_cell)
                    self._append_inline_content(paragraph, html_cell, force_bold=html_cell.name.lower() == "th")
                    styles = self._styles(html_cell)
                    if styles.get("background-color"):
                        self._set_cell_shading(cell, styles["background-color"])
            return
        if name == "hr":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "808080")
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
            return
        for child in node.children:
            if isinstance(child, Tag):
                self._append_html_node(document, child)

    def _append_inline_content(
        self,
        paragraph,
        node: Tag,
        inherited: dict | None = None,
        heading_level: int | None = None,
        force_bold: bool = False,
    ) -> None:
        inherited = dict(inherited or {})
        name = node.name.lower()
        styles = self._styles(node)
        state = {**inherited, **styles}
        if name in {"strong", "b"} or force_bold:
            state["font-weight"] = "bold"
        if name in {"em", "i"}:
            state["font-style"] = "italic"
        if name == "u":
            state["text-decoration"] = "underline"
        if name in {"s", "strike", "del"}:
            state["text-decoration"] = "line-through"
        if name == "sup":
            state["vertical-align"] = "super"
        if name == "sub":
            state["vertical-align"] = "sub"

        for child in node.children:
            if isinstance(child, NavigableString):
                if not str(child):
                    continue
                run = paragraph.add_run(str(child))
                self._apply_run_style(run, state, heading_level)
            elif isinstance(child, Tag):
                if child.name.lower() == "br":
                    paragraph.add_run().add_break()
                else:
                    self._append_inline_content(paragraph, child, state, heading_level, force_bold)

    def _apply_paragraph_style(self, paragraph, node: Tag) -> None:
        styles = self._styles(node)
        alignment = styles.get("text-align") or node.attrs.get("align")
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        }.get(str(alignment).lower())
        line_height = styles.get("line-height")
        if line_height:
            try:
                paragraph.paragraph_format.line_spacing = float(line_height)
            except ValueError:
                value = self._css_length_pt(line_height)
                if value is not None:
                    paragraph.paragraph_format.line_spacing = Pt(value)
        for css_name, attr in (
            ("margin-top", "space_before"),
            ("margin-bottom", "space_after"),
            ("margin-left", "left_indent"),
            ("margin-right", "right_indent"),
            ("text-indent", "first_line_indent"),
        ):
            value = self._css_length_pt(styles.get(css_name))
            if value is not None:
                setattr(paragraph.paragraph_format, attr, Pt(value))

    def _apply_run_style(self, run, styles: dict, heading_level: int | None = None) -> None:
        weight = str(styles.get("font-weight", "")).lower()
        run.bold = weight in {"bold", "bolder", "600", "700", "800", "900"} or bool(heading_level)
        run.italic = str(styles.get("font-style", "")).lower() == "italic"
        decoration = str(styles.get("text-decoration", "")).lower()
        run.underline = "underline" in decoration
        run.font.strike = "line-through" in decoration
        vertical = str(styles.get("vertical-align", "")).lower()
        run.font.superscript = vertical in {"super", "sup"}
        run.font.subscript = vertical == "sub"

        size = self._css_length_pt(styles.get("font-size"))
        if size is None and heading_level:
            size = {1: 24, 2: 20, 3: 18, 4: 16, 5: 14, 6: 12}.get(heading_level)
        if size is not None:
            run.font.size = Pt(size)

        family = str(styles.get("font-family", "")).split(",")[0].strip(" '\"")
        if family:
            run.font.name = family
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), family)
        color = self._normalise_color(styles.get("color"))
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        background = self._normalise_color(styles.get("background-color") or styles.get("background"))
        if background:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), background)
            run._element.get_or_add_rPr().append(shading)

    @staticmethod
    def _styles(node: Tag) -> dict[str, str]:
        result: dict[str, str] = {}
        for declaration in str(node.attrs.get("style", "")).split(";"):
            if ":" not in declaration:
                continue
            key, value = declaration.split(":", 1)
            result[key.strip().lower()] = value.strip()
        return result

    @staticmethod
    def _css_length_pt(value: str | None) -> float | None:
        if not value:
            return None
        text = str(value).strip().lower()
        try:
            if text.endswith("pt"):
                return float(text[:-2])
            if text.endswith("px"):
                return float(text[:-2]) * 0.75
            if text.endswith("cm"):
                return float(text[:-2]) * 28.3464567
            if text.endswith("mm"):
                return float(text[:-2]) * 2.83464567
            if text.endswith("em"):
                return float(text[:-2]) * 12
            return float(text)
        except ValueError:
            return None

    @staticmethod
    def _normalise_color(value: str | None) -> str | None:
        if not value:
            return None
        text = str(value).strip().lower()
        if text.startswith("#") and len(text) in {4, 7}:
            raw = text[1:]
            if len(raw) == 3:
                raw = "".join(char * 2 for char in raw)
            return raw.upper()
        if text.startswith("rgb(") and text.endswith(")"):
            try:
                parts = [int(part.strip()) for part in text[4:-1].split(",")]
                return "".join(f"{max(0, min(255, part)):02X}" for part in parts[:3])
            except (ValueError, TypeError):
                return None
        return None

    def _set_cell_shading(self, cell, color: str) -> None:
        fill = self._normalise_color(color)
        if not fill:
            return
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    def _render_content(self, content: ReportContent) -> str:
        sections = "".join(self._render_section(section) for section in content.sections)
        return f"<h1>{escape(content.title)}</h1>{sections}"

    def _render_section(self, section: ReportSection) -> str:
        content = section.content or ""
        if section.type == "html":
            section_body = content
        else:
            section_body = f"<p>{escape(content)}</p>" if content else ""
        source = ""
        if section.source:
            source = f"<p class=\"source\">来源：{escape(', '.join(section.source))}</p>"
        ai_badge = "<span class=\"badge\">AI 草稿</span>" if section.ai_generated else ""
        return f"<section><h2>{escape(section.title)}{ai_badge}</h2>{section_body}{source}</section>"

    def _html_document(self, title: str, body: str, editor_config: ReportEditorConfig | None = None) -> str:
        page = editor_config.page if editor_config else None
        orientation = page.orientation if page else "portrait"
        if page and page.size and page.size.width and page.size.height:
            width, height = page.size.width, page.size.height
            if orientation == "landscape":
                width, height = max(width, height), min(width, height)
            else:
                width, height = min(width, height), max(width, height)
            page_size = f"{width}cm {height}cm"
        else:
            page_size = f"A4 {orientation}"
        margins = page.margin if page else None
        page_margin = (
            f"{margins.top}cm {margins.right}cm {margins.bottom}cm {margins.left}cm"
            if margins
            else "2.54cm"
        )
        main_width = "1180px" if orientation == "landscape" else "920px"
        return f"""<!doctype html>
<html lang=\"zh-CN\">
<head>
  <meta charset=\"utf-8\" />
  <title>{escape(title)}</title>
  <style>
    @page {{ size: {page_size}; margin: {page_margin}; }}
    body {{ margin: 0; background: #f5f7fb; color: #1f2a3d; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Microsoft YaHei', sans-serif; }}
    main {{ max-width: {main_width}; margin: 32px auto; padding: 48px 64px; background: #fff; border-radius: 18px; box-shadow: 0 18px 50px rgba(34, 75, 130, .12); }}
    h1 {{ margin: 0 0 28px; text-align: center; font-size: 28px; }}
    h2 {{ margin: 28px 0 12px; font-size: 20px; color: #123b78; }}
    p {{ line-height: 1.9; font-size: 16px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0; }}
    th, td {{ border: 1px solid #d8e2ef; padding: 9px 11px; text-align: left; }}
    th {{ background: #eef5ff; }}
    .badge {{ margin-left: 8px; font-size: 12px; color: #2878ff; background: #eaf3ff; border-radius: 999px; padding: 2px 8px; }}
    .source, .empty {{ color: #7b8aa3; font-size: 13px; }}
  </style>
</head>
<body><main>{body}</main></body>
</html>"""
