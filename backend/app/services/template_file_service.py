from __future__ import annotations

import io
import re
import zipfile
from html import escape
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fastapi import HTTPException, UploadFile, status

ALLOWED_WORD_SUFFIXES = {".docx"}
MAX_WORD_FILE_SIZE = 20 * 1024 * 1024
MAX_ZIP_ENTRIES = 2000
MAX_ZIP_ENTRY_SIZE = 50 * 1024 * 1024
MAX_ZIP_TOTAL_SIZE = 200 * 1024 * 1024
DOCX_REQUIRED_ENTRIES = {"[Content_Types].xml", "word/document.xml"}


class TemplateFileService:
    def __init__(self, storage_dir: Path | None = None):
        self.storage_dir = storage_dir or Path(__file__).resolve().parents[2] / "storage" / "templates"

    async def save_word(self, upload: UploadFile) -> dict[str, str | int]:
        original_filename = Path(upload.filename or "").name
        suffix = Path(original_filename).suffix.lower()
        if suffix not in ALLOWED_WORD_SUFFIXES:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="仅支持 .docx 格式；旧版 .doc 无法安全解析，请另存为 .docx 后上传",
            )

        content = await upload.read(MAX_WORD_FILE_SIZE + 1)
        if not content:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="上传文件不能为空")
        if len(content) > MAX_WORD_FILE_SIZE:
            raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="Word 文件不能超过 20MB")
        self._validate_docx(content)

        self.storage_dir.mkdir(parents=True, exist_ok=True)
        safe_stem = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", Path(original_filename).stem).strip("_") or "template"
        stored_name = f"{uuid4().hex}_{safe_stem[:80]}{suffix}"
        target = self.storage_dir / stored_name
        try:
            target.write_bytes(content)
        except OSError as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="模板文件保存失败") from exc

        return {
            "original_filename": original_filename,
            "file_path": str(target.resolve()),
            "file_size": len(content),
            "mime_type": upload.content_type or "application/octet-stream",
        }

    @staticmethod
    def _validate_docx(content: bytes) -> None:
        if not content.startswith(b"PK"):
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="文件不是有效的 DOCX（ZIP 签名错误）")
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > MAX_ZIP_ENTRIES:
                    raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="DOCX 内部条目过多")
                names = {entry.filename for entry in entries}
                if not DOCX_REQUIRED_ENTRIES.issubset(names):
                    raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="DOCX 缺少必需文档条目")
                total = 0
                for entry in entries:
                    if entry.file_size > MAX_ZIP_ENTRY_SIZE:
                        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="DOCX 单个解压条目过大")
                    total += entry.file_size
                    if total > MAX_ZIP_TOTAL_SIZE:
                        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="DOCX 解压后总体积过大")
                    if entry.compress_size > 0 and entry.file_size / entry.compress_size > 1000:
                        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="DOCX 压缩比异常")
        except zipfile.BadZipFile as exc:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="文件不是有效的 DOCX 压缩包") from exc

    def delete(self, file_path: str | None) -> None:
        if not file_path:
            return
        path = self._resolve_storage_path(file_path)
        try:
            path.unlink(missing_ok=True)
        except OSError:
            pass

    def read(self, file_path: str | None) -> bytes:
        path = self._require_file(file_path)
        try:
            content = path.read_bytes()
        except OSError as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="模板文件读取失败") from exc
        if len(content) > MAX_WORD_FILE_SIZE:
            raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="Word 文件不能超过 20MB")
        self._validate_docx(content)
        return content

    def extract_html(self, file_path: str | None) -> str:
        path = self._require_file(file_path)
        if path.suffix.lower() != ".docx":
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="编辑器加载内容暂只支持 .docx 格式，请将旧版 .doc 文件另存为 .docx 后重新上传",
            )
        # Re-read and validate historical files before handing them to
        # python-docx so old database paths cannot bypass ZIP resource limits.
        content = self.read(file_path)
        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Word 模板内容解析失败") from exc

        paragraphs = {id(item._p): item for item in document.paragraphs}
        tables = {id(item._tbl): item for item in document.tables}
        blocks: list[str] = []
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = paragraphs.get(id(child))
                if paragraph is not None:
                    blocks.append(TemplateFileService._paragraph_html(paragraph))
            elif child.tag.endswith("}tbl"):
                table = tables.get(id(child))
                if table is not None:
                    blocks.append(TemplateFileService._table_html(table))
        html = "".join(blocks).strip()
        return html or "<p></p>"

    @staticmethod
    def _paragraph_html(paragraph) -> str:
        runs = []
        for run in paragraph.runs:
            text = escape(run.text).replace("\t", "&emsp;").replace("\n", "<br>")
            if not text:
                continue
            styles = TemplateFileService._run_styles(run)
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            if styles:
                text = f'<span style="{styles}">{text}</span>'
            runs.append(text)

        content = "".join(runs) or "<br>"
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        heading_match = re.search(r"heading\s*([1-6])|标题\s*([1-6])", style_name)
        tag = f"h{heading_match.group(1) or heading_match.group(2)}" if heading_match else "p"
        paragraph_styles = TemplateFileService._paragraph_styles(paragraph)
        style_attr = f' style="{paragraph_styles}"' if paragraph_styles else ""
        return f"<{tag}{style_attr}>{content}</{tag}>"

    @staticmethod
    def _run_styles(run) -> str:
        font = run.font
        styles: list[str] = []
        font_name = font.name or TemplateFileService._east_asia_font(run)
        if font_name:
            styles.append(f"font-family:{TemplateFileService._css_font(font_name)}")
        if font.size:
            styles.append(f"font-size:{font.size.pt:.2f}pt")
        if font.color and font.color.type is not None and font.color.rgb:
            styles.append(f"color:#{font.color.rgb}")
        try:
            highlight_color = font.highlight_color
        except ValueError:
            # 部分 Word/WPS 文档会写入 python-docx 未映射的 w:highlight="none"。
            highlight_color = None
        if highlight_color is not None:
            highlight = TemplateFileService._highlight_color(highlight_color)
            if highlight:
                styles.append(f"background-color:{highlight}")
        if font.strike:
            styles.append("text-decoration:line-through")
        if font.superscript:
            styles.append("vertical-align:super;font-size:smaller")
        elif font.subscript:
            styles.append("vertical-align:sub;font-size:smaller")
        return ";".join(styles)

    @staticmethod
    def _paragraph_styles(paragraph) -> str:
        fmt = paragraph.paragraph_format
        styles: list[str] = []
        alignment = paragraph.alignment
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
        }
        if alignment in alignment_map:
            styles.append(f"text-align:{alignment_map[alignment]}")
        if fmt.left_indent:
            styles.append(f"margin-left:{fmt.left_indent.cm:.2f}cm")
        if fmt.right_indent:
            styles.append(f"margin-right:{fmt.right_indent.cm:.2f}cm")
        if fmt.first_line_indent:
            styles.append(f"text-indent:{fmt.first_line_indent.cm:.2f}cm")
        if fmt.space_before:
            styles.append(f"margin-top:{fmt.space_before.pt:.2f}pt")
        if fmt.space_after:
            styles.append(f"margin-bottom:{fmt.space_after.pt:.2f}pt")
        if fmt.line_spacing:
            if hasattr(fmt.line_spacing, "pt"):
                styles.append(f"line-height:{fmt.line_spacing.pt:.2f}pt")
            elif isinstance(fmt.line_spacing, (int, float)):
                styles.append(f"line-height:{float(fmt.line_spacing):.2f}")
        if fmt.keep_with_next:
            styles.append("break-after:avoid")
        if fmt.page_break_before:
            styles.append("break-before:page")
        return ";".join(styles)

    @staticmethod
    def _table_html(table) -> str:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_styles = ["border:1px solid #000", "padding:4px 6px"]
                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                    cell_styles.append("vertical-align:middle")
                elif cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
                    cell_styles.append("vertical-align:bottom")
                width = cell.width
                if width:
                    cell_styles.append(f"width:{width.cm:.2f}cm")
                content = "".join(TemplateFileService._paragraph_html(item) for item in cell.paragraphs)
                cells.append(f'<td style="{";".join(cell_styles)}">{content}</td>')
            rows.append(f"<tr>{''.join(cells)}</tr>")
        return f'<table style="width:100%;border-collapse:collapse;table-layout:fixed"><tbody>{"".join(rows)}</tbody></table>'

    @staticmethod
    def _east_asia_font(run) -> str | None:
        r_pr = run._element.rPr
        if r_pr is None or r_pr.rFonts is None:
            return None
        return r_pr.rFonts.get(qn("w:eastAsia")) or r_pr.rFonts.get(qn("w:ascii"))

    @staticmethod
    def _css_font(value: str) -> str:
        safe = value.replace('"', "").replace("'", "").strip()
        return f"'{safe}', sans-serif"

    @staticmethod
    def _highlight_color(value) -> str | None:
        colors = {
            "YELLOW (7)": "#ffff00",
            "BRIGHT_GREEN (4)": "#00ff00",
            "TURQUOISE (3)": "#00ffff",
            "PINK (5)": "#ff00ff",
            "BLUE (2)": "#0000ff",
            "RED (6)": "#ff0000",
            "DARK_BLUE (9)": "#000080",
            "TEAL (10)": "#008080",
            "GREEN (11)": "#008000",
            "VIOLET (12)": "#800080",
            "DARK_RED (13)": "#800000",
            "DARK_YELLOW (14)": "#808000",
            "GRAY_50 (15)": "#808080",
            "GRAY_25 (16)": "#c0c0c0",
            "BLACK (1)": "#000000",
        }
        return colors.get(str(value))

    def _resolve_storage_path(self, file_path: str) -> Path:
        root = self.storage_dir.resolve()
        candidate = Path(file_path)
        if not candidate.is_absolute():
            candidate = root / candidate
        resolved = candidate.resolve()
        try:
            resolved.relative_to(root)
        except ValueError as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="模板文件路径超出存储目录",
            ) from exc
        return resolved

    def _require_file(self, file_path: str | None) -> Path:
        if not file_path:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="该模板没有关联 Word 文件")
        path = self._resolve_storage_path(file_path)
        if not path.is_file():
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="模板文件不存在")
        return path
